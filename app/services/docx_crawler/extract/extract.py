import json
import os
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

from app.services.docx_crawler.extract.config import (
    TOTALS_ROW_KEYWORDS,
    TOTALS_ROW_MIN_MATCHES,
    SAE_KEYWORDS,
)


def parse_table_title(title_text: str) -> tuple[str | None, str]:
    """
    Extracts the table number and description from a title string.

      "Table 2.1.1 Serious AEs by Preferred term"
      -> ("2.1.1", "Serious AEs by Preferred term")
    """
    pattern = r"^[Tt]able\s+(\d+(?:\.\d+)*)[\s.:]*(.*)$"
    match = re.match(pattern, title_text.strip())
    if match:
        return match.group(1), match.group(2).strip()
    return None, title_text.strip()


def iter_block_items(document: Document):
    """
    Walk the document body in order, yielding each Paragraph or Table as we
    encounter it. python-docx only gives document.paragraphs OR document.tables
    separately, with no indication of how they interleave; this restores order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def is_sae_title(text: str) -> bool:
    """
    Check if a paragraph's text looks like an SAE table title.

    FIX: leading word-boundary instead of plain substring. This avoids a keyword
    accidentally matching inside an unrelated word, while still matching plurals
    such as "SAEs" (\\bsae matches the start of "saes").
    """
    lowered = text.lower()
    return any(re.search(rf"\b{re.escape(keyword)}", lowered) for keyword in SAE_KEYWORDS)


def is_totals_row(first_cell_text: str) -> bool:
    """Check if a table row is the 'totals' row, based on its first cell."""
    lowered = first_cell_text.lower()
    matched = sum(1 for word in TOTALS_ROW_KEYWORDS if word in lowered)
    return matched >= TOTALS_ROW_MIN_MATCHES


def first_int(text: str) -> int:
    """
    Pull the leading integer out of a cell, tolerating thousands separators and
    trailing text. FIX: replaces the old `str.isdigit()` check, which silently
    returned 0 for values like '19 (100%)' or '1,234'.
    """
    match = re.search(r"\d[\d,]*", text or "")
    return int(match.group(0).replace(",", "")) if match else 0


def parse_cell_value(val_str: str) -> tuple[str, str]:
    """
    Splits a cell like '1 (33%)' into ('1', '33%'), or '0'/'' into ('0', '0%').

    FIX: strip spaces inside the bracket so '1 (33 %)' -> ('1', '33%').
    NOTE/assumption: percentages only ever come from brackets (per the spec, we
    extract rather than calculate). A bare non-zero count therefore has no
    percentage to report; in the sample data bare cells are always '0'.
    """
    val_str = (val_str or "").strip()
    if not val_str or val_str == "0":
        return "0", "0%"

    if "(" in val_str:
        count, _, rest = val_str.partition("(")
        percent = rest.replace(")", "").replace(" ", "").strip()
        return count.strip(), percent

    return val_str, "0%"


def find_sae_tables(document: Document):
    """
    Walk the document, tracking the most recent non-empty paragraph as a
    candidate title. When we hit a table, if that candidate matches an SAE
    keyword we keep the table. The title is reset after every table so it
    can't accidentally attach to a second, unrelated table further down.
    """
    candidate_title = None
    matches = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                candidate_title = text
            # blank paragraph -> leave candidate_title as-is, so a blank line
            # between the title and the table doesn't break the pairing
        elif isinstance(block, Table):
            if candidate_title and is_sae_title(candidate_title):
                table_num, table_name = parse_table_title(candidate_title)
                matches.append(
                    {
                        "title": candidate_title,
                        "table_number": table_num,
                        "table_name": table_name,
                        "table": block,
                    }
                )
            candidate_title = None

    return matches


def build_cell(compound_name: str, cell_text: str) -> dict:
    """
    "1 (33%)"  -> count=1, int_percent=33, percent="33%"
    "0" or ""  -> count=0, int_percent=0,  percent=None
    "3 (16 %)" -> count=3, int_percent=16, percent="16%"  (internal space dropped)
    """
    text = (cell_text or "").strip()
    if not text or text == "0":
        return {"compound": compound_name, "count": 0, "int_percent": 0, "percent": None}
    if "(" in text:
        count_part, _, rest = text.partition("(")
        percent = rest.replace(")", "").replace(" ", "").strip()   # "33%"
        num = percent.rstrip("%")
        int_percent = round(float(num)) if num else 0
        return {"compound": compound_name, "count": first_int(count_part),
                "int_percent": int_percent, "percent": percent}
    # non-zero count, no bracket: count present, but no percentage to extract
    return {"compound": compound_name, "count": first_int(text), "int_percent": 0, "percent": None}


def extract_found_tables_into_json(found: list) -> dict:
    extracted_data = {"tables": []}

    for match in found:
        title = match.get("table_name") or match.get("title", "")
        table_data = {
            "table_number": match.get("table_number", ""),
            "table_title": title,
            "term_data": [],
            "total_with_SAE": [],
        }

        table = match["table"]
        row_to_term, terms_list, totals_row = {}, [], None

        # locate the header row + the term column
        header_idx, term_col = None, 0
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "preferred term" in cell.text.strip().lower():
                    header_idx, term_col = r_idx, c_idx
                    break
            if header_idx is not None:
                break
        if header_idx is None:  # fallback: first non-empty row is the header
            for r_idx, row in enumerate(table.rows):
                if any(cell.text.strip() for cell in row.cells):
                    header_idx, term_col = r_idx, 0
                    break
        if header_idx is None:
            continue  # genuinely empty table

        # TERM COLUMN: walk DOWN the term column to find term rows + the totals row
        term_column_cells = [cell.text.strip() for cell in table.columns[term_col].cells]
        for row_pos, label in enumerate(term_column_cells):
            if row_pos == header_idx or not label:
                continue
            if is_totals_row(label):
                totals_row = row_pos
            else:
                term_entry = {"term": label, "compounds": []}
                terms_list.append(term_entry)
                row_to_term[row_pos] = term_entry

        # COMPOUND COLUMNS: every column except the term column is one compound
        for col_idx in range(len(table.columns)):
            if col_idx == term_col:
                continue
            cells = [cell.text.strip() for cell in table.columns[col_idx].cells]
            if header_idx >= len(cells) or not cells[header_idx]:
                continue
            compound_name = cells[header_idx]
            for row_pos, cell_text in enumerate(cells):
                if row_pos == header_idx:
                    continue
                if totals_row is not None and row_pos == totals_row:
                    table_data["total_with_SAE"].append(
                        {"compound": compound_name, "count": first_int(cell_text)}
                    )
                elif row_pos in row_to_term:
                    row_to_term[row_pos]["compounds"].append(
                        build_cell(compound_name, cell_text)
                    )

        table_data["term_data"] = terms_list
        extracted_data["tables"].append(table_data)

    return extracted_data


# Top level function:
def parse_docx_file(file_path: str):
    doc = Document(file_path)
    found = find_sae_tables(doc)
    return extract_found_tables_into_json(found)
