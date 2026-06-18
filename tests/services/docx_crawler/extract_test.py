from pathlib import Path

from docx import Document
from app.services.docx_crawler.extract.extract import (
    parse_table_title,
    is_sae_title,
    is_totals_row,
    first_int,
    build_cell,
    locate_header,
    compound_columns,
    extract_table,
    find_sae_tables,
    extract_found_tables_into_json,
)


RESOURCE = Path(__file__).parent / "resources" / "client1.docx"


# ---------------------------------------------------------------------------
# parse_table_title
# ---------------------------------------------------------------------------


def test_parse_table_title_number_and_description():
    assert parse_table_title("Table 2.1.1 Summary of Serious Adverse Events") == (
        "2.1.1",
        "Summary of Serious Adverse Events",
    )


def test_parse_table_title_multi_segment_number_with_colon():
    assert parse_table_title("Table 14.3.1.2: Serious Adverse Events") == (
        "14.3.1.2",
        "Serious Adverse Events",
    )


def test_parse_table_title_no_number_returns_none_and_full_text():
    assert parse_table_title("Serious Adverse Events") == (
        None,
        "Serious Adverse Events",
    )


def test_parse_table_title_non_numeric_label_is_not_captured():
    # 'A.1' is not \d+(\.\d+)* so the whole string falls through to the description
    assert parse_table_title("Table A.1 SAE Summary") == (None, "Table A.1 SAE Summary")


# ---------------------------------------------------------------------------
# is_sae_title
# ---------------------------------------------------------------------------


def test_is_sae_title_matches_serious():
    assert is_sae_title("Summary of Serious Adverse Events") is True


def test_is_sae_title_matches_sae_acronym():
    assert is_sae_title("Table 1 SAE overview") is True


def test_is_sae_title_matches_plural_saes():
    assert is_sae_title("Summary of SAEs") is True


def test_is_sae_title_rejects_non_sae_table():
    assert is_sae_title("Adverse Events Leading to Death") is False


# ---------------------------------------------------------------------------
# is_totals_row
# ---------------------------------------------------------------------------


def test_is_totals_row_matches_sae_totals_label():
    assert is_totals_row("Total number of participants with SAE") is True


def test_is_totals_row_matches_sae_number_label():
    assert is_totals_row("Number of participants with any SAE") is True


def test_is_totals_row_rejects_term_row():
    assert is_totals_row("Seizure") is False


# ---------------------------------------------------------------------------
# first_int
# ---------------------------------------------------------------------------


def test_first_int_plain():
    assert first_int("3") == 3


def test_first_int_strips_percentage_suffix():
    assert first_int("19 (100%)") == 19


def test_first_int_handles_thousands_separator():
    assert first_int("1,234") == 1234


def test_first_int_empty_is_zero():
    assert first_int("") == 0


# ---------------------------------------------------------------------------
# build_cell
# ---------------------------------------------------------------------------


def test_build_cell_count_and_percent():
    assert build_cell("Placebo", "1 (33%)") == {
        "compound": "Placebo",
        "count": 1,
        "percent": 33,
        "percent_str": "33%",
    }


def test_build_cell_bare_zero_has_null_percent():
    assert build_cell("Placebo", "0") == {
        "compound": "Placebo",
        "count": 0,
        "percent": 0,
        "percent_str": None,
    }


def test_build_cell_empty_string_has_null_percent():
    assert build_cell("Placebo", "") == {
        "compound": "Placebo",
        "count": 0,
        "percent": 0,
        "percent_str": None,
    }


def test_build_cell_strips_internal_space_in_percent():
    assert build_cell("Compound X", "3 (16 %)") == {
        "compound": "Compound X",
        "count": 3,
        "percent": 16,
        "percent_str": "16%",
    }


def test_build_cell_decimal_percent_rounds_up_int_keeps_display():
    # percent keeps the decimal value
    assert build_cell("Compound X", "2 (11.5%)") == {
        "compound": "Compound X",
        "count": 2,
        "percent": 11.5,
        "percent_str": "11.5%",
    }


def test_build_cell_non_zero_without_bracket_has_null_percent():
    # shouldn't occur in real data, but documents the handled fallback
    assert build_cell("Placebo", "5") == {
        "compound": "Placebo",
        "count": 5,
        "percent": 0,
        "percent_str": None,
    }


# ---------------------------------------------------------------------------
# locate_header / compound_columns
# ---------------------------------------------------------------------------


# =========  HELPER ===========
def _table_from_rows(rows: list[list[str]]):
    """Build a throwaway docx table from a list of rows for header/column tests."""
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    return table


# =============================


def test_locate_header_finds_preferred_term_cell():
    table = _table_from_rows(
        [
            ["Preferred Term", "Placebo", "Compound X"],
            ["Seizure", "1 (33%)", "3 (16%)"],
        ]
    )
    assert locate_header(table) == (0, 0)


def test_locate_header_falls_back_to_first_non_empty_row():
    table = _table_from_rows(
        [
            ["Adverse Event", "Placebo", "Compound X"],
            ["Seizure", "1 (33%)", "3 (16%)"],
        ]
    )
    assert locate_header(table) == (0, 0)


def test_compound_columns_lists_non_term_columns_in_order():
    table = _table_from_rows(
        [
            ["Preferred Term", "Placebo", "Compound X"],
            ["Seizure", "1 (33%)", "3 (16%)"],
        ]
    )
    assert compound_columns(table, header_idx=0, term_col=0) == [
        (1, "Placebo"),
        (2, "Compound X"),
    ]


# ---------------------------------------------------------------------------
# extract_table  (built table)
# ---------------------------------------------------------------------------


def test_extract_table_builds_term_first_structure():
    table = _table_from_rows(
        [
            ["Preferred Term", "Placebo", "Compound X"],
            ["Seizure", "1 (33%)", "3 (16%)"],
            ["Nausea", "0", "2 (11%)"],
            ["Total number of participants with SAE", "3", "19"],
        ]
    )
    result = extract_table(table, "2.1.1", "Summary of Serious Adverse Events")
    assert result["table_number"] == "2.1.1"
    assert result["total_with_SAE"] == [
        {"compound": "Placebo", "count": 3},
        {"compound": "Compound X", "count": 19},
    ]
    assert result["term_data"][0] == {
        "term": "Seizure",
        "compounds": [
            {"compound": "Placebo", "count": 1, "percent": 33, "percent_str": "33%"},
            {"compound": "Compound X", "count": 3, "percent": 16, "percent_str": "16%"},
        ],
    }
    assert [t["term"] for t in result["term_data"]] == ["Seizure", "Nausea"]


def test_extract_table_returns_none_for_empty_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)  # single empty cell
    assert extract_table(table, "x", "x") is None


# ---------------------------------------------------------------------------
# find_sae_tables / full pipeline  (real fixture)
# ---------------------------------------------------------------------------


def test_find_sae_tables_keeps_only_sae_titled_tables():
    doc = Document(str(RESOURCE))
    matches = find_sae_tables(doc)
    titles = [m["title"] for m in matches]
    assert len(matches) == 1
    assert "Serious" in titles[0]
    assert all("Death" not in t for t in titles)


def test_full_pipeline_on_fixture():
    doc = Document(str(RESOURCE))
    out = extract_found_tables_into_json(find_sae_tables(doc))
    assert len(out["tables"]) == 1
    table = out["tables"][0]
    assert table["table_number"] == "2.1.1"
    assert [t["term"] for t in table["term_data"]] == [
        "Seizure",
        "Nausea",
        "Headache",
        "Bleeding",
    ]
    assert table["total_with_SAE"] == [
        {"compound": "Placebo", "count": 3},
        {"compound": "Compound X", "count": 19},
    ]
